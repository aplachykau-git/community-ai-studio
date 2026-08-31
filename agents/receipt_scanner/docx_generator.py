"""
DOCX Expense Report Generator for Receipt Scanner Agent.
Generates styled Microsoft Word (.docx) reimbursement reports populated from
receipt vision scans and template assets.
"""

import base64
import copy
import datetime
import io
import os
import re
from typing import Any, Dict, List, Optional

try:
    import docx
    from docx import Document
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import nsdecls, qn
    from docx.shared import Inches, Pt, RGBColor
    from PIL import Image, ImageOps
except ImportError:
    docx = None
    Document = None


def calculate_receipt_totals(receipts_data: List[Dict[str, Any]], target_currency: str = "USD") -> tuple[str, str]:
    """Calculates total original currency and total target currency amounts."""
    total_target = 0.0
    orig_by_curr: Dict[str, float] = {}

    for r in receipts_data or []:
        # Original sum
        raw_sc = r.get("sum_curr") or r.get("sum_pln") or "0"
        m_sc = re.search(r"([\d.,]+)\s*([A-Za-z]+)?", str(raw_sc))
        if m_sc:
            try:
                amt = float(m_sc.group(1).replace(",", "."))
                curr = m_sc.group(2) or "PLN"
                orig_by_curr[curr.upper()] = orig_by_curr.get(curr.upper(), 0.0) + amt
            except ValueError:
                pass

        # Target sum
        raw_st = r.get("sum_target") or r.get("sum_usd") or "0"
        m_st = re.search(r"([\d.,]+)", str(raw_st))
        if m_st:
            try:
                amt = float(m_st.group(1).replace(",", "."))
                total_target += amt
            except ValueError:
                pass

    if orig_by_curr:
        orig_strs = [f"{amt:.2f} {curr}" for curr, amt in orig_by_curr.items()]
        orig_tot_str = " + ".join(orig_strs)
    else:
        orig_tot_str = "0.00 PLN"

    target_tot_str = f"{total_target:.2f} {target_currency.upper()}"
    return orig_tot_str, target_tot_str


def _set_row_cells(row, item_num, cat, desc, sc, st):
    """Sets standard expense item row values and formatting."""
    row.cells[0].paragraphs[0].text = str(item_num)
    row.cells[1].paragraphs[0].text = str(cat or "")
    row.cells[2].paragraphs[0].text = str(desc or "")
    row.cells[3].paragraphs[0].text = str(sc or "")
    row.cells[4].paragraphs[0].text = str(st or "")

    for idx, cell in enumerate(row.cells):
        p = cell.paragraphs[0]
        p.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
            if idx in [0, 1]
            else (WD_ALIGN_PARAGRAPH.RIGHT if idx in [3, 4] else WD_ALIGN_PARAGRAPH.LEFT)
        )
        for r in p.runs:
            r.font.name = "Arial"
            r.font.size = Pt(9.5)


def generate_expense_report_docx(
    title: str,
    receipts_data: List[Dict[str, Any]],
    target_currency: str = "USD",
    exchange_rate: float = 1.0,
    rate_source: str = "Official Banking Exchange Rates",
    rate_url: str = "",
    approved_budget: Optional[str] = None,
    community_name: str = "Community",
    session_images: Optional[List[Any]] = None,
    output_path: Optional[str] = None,
    template_path: Optional[str] = None,
) -> bytes:
    """Generates a styled .docx expense report document from the 1-to-1 template."""
    if Document is None:
        raise RuntimeError("python-docx is not installed in the current environment.")

    use_existing_template = bool(template_path and os.path.exists(template_path))
    if not use_existing_template:
        default_tpl = os.path.join(os.path.dirname(__file__), "assets", "expense_report_template.docx")
        if os.path.exists(default_tpl):
            template_path = default_tpl
            use_existing_template = True

    if use_existing_template:
        doc = Document(template_path)
    else:
        doc = Document()

    today_str = datetime.date.today().strftime("%d.%m.%Y")
    total_orig_str, total_target_str = calculate_receipt_totals(receipts_data, target_currency)

    # 1. Replace paragraph tags
    rate_bank_text = rate_url if rate_url else (f"({rate_source})" if rate_source and rate_source != "N/A" else "")
    for p in doc.paragraphs:
        if "{{TITLE}}" in p.text:
            p.text = p.text.replace("{{TITLE}}", title or "Expense Reimbursement Report")
            if p.runs:
                p.runs[0].font.name = "Arial"
                p.runs[0].font.size = Pt(18)
                p.runs[0].font.bold = True

        if "{{Current date}}" in p.text or "{{Bank link}}" in p.text:
            p.text = p.text.replace("{{Current date}}", today_str).replace("{{Bank link}}", rate_bank_text).strip()
            for r in p.runs:
                r.font.name = "Arial"
                r.font.size = Pt(9.5)
                r.font.italic = True

        if "{{EUR/USD}}" in p.text or "{{Current exchange rate}}" in p.text:
            rate_str = f"{exchange_rate:.4f}" if exchange_rate else "1.0000"
            p.text = (
                p.text.replace("{{EUR/USD}}", target_currency.upper())
                .replace("{{Current exchange rate}}", rate_str)
                .replace("{{Local Currency Code}}", "PLN")
            )
            for r in p.runs:
                r.font.name = "Arial"
                r.font.size = Pt(9.5)
                r.font.bold = True

    # 2. Populate Expenses Table (Table 0)
    if doc.tables:
        exp_table = doc.tables[0]
        n_items = len(receipts_data or [])

        if n_items > 0 and len(exp_table.rows) >= 4:
            # Row 1 is template
            item0 = receipts_data[0]
            _set_row_cells(
                exp_table.rows[1],
                1,
                item0.get("category", ""),
                item0.get("desc", ""),
                item0.get("sum_curr", ""),
                item0.get("sum_target", ""),
            )

            # Insert additional items
            for i in range(1, n_items):
                item = receipts_data[i]
                new_tr = copy.deepcopy(exp_table.rows[1]._tr)
                total_tr = exp_table.rows[-2]._tr
                total_tr.addprevious(new_tr)

            for i in range(1, n_items):
                item = receipts_data[i]
                _set_row_cells(
                    exp_table.rows[i + 1],
                    i + 1,
                    item.get("category", ""),
                    item.get("desc", ""),
                    item.get("sum_curr", ""),
                    item.get("sum_target", ""),
                )

            # Update Total Spent row
            tot_row = exp_table.rows[-2]
            tot_row.cells[3].paragraphs[0].text = total_orig_str
            tot_row.cells[4].paragraphs[0].text = total_target_str
            for idx in [3, 4]:
                p = tot_row.cells[idx].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(9.5)
                    r.font.bold = True

            # Update Approved Budget row
            bud_row = exp_table.rows[-1]
            bud_str = approved_budget or f"150.00 {target_currency.upper()}"
            bud_row.cells[4].paragraphs[0].text = bud_str
            p = bud_row.cells[4].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for r in p.runs:
                r.font.name = "Arial"
                r.font.size = Pt(9.5)
                r.font.bold = True

    # 3. Collect Proof Images
    images_to_embed: List[bytes] = []

    def _normalize_image_bytes(raw: Any) -> Optional[bytes]:
        if not raw:
            return None
        if isinstance(raw, bytes):
            return raw
        if isinstance(raw, str):
            # Check if existing local file path
            if os.path.exists(raw):
                try:
                    with open(raw, "rb") as f:
                        return f.read()
                except Exception as err:
                    print(f"⚠️ [Image Read Error] Could not read {raw}: {err}")
            # Check if base64 data URL
            if raw.startswith("data:") and "," in raw:
                try:
                    return base64.b64decode(raw.split(",", 1)[1])
                except Exception:
                    pass
            # Try raw base64 decode
            try:
                decoded = base64.b64decode(raw)
                if len(decoded) > 10:
                    return decoded
            except Exception:
                pass
        return None

    if session_images:
        for s_img in session_images:
            norm_b = _normalize_image_bytes(s_img)
            if norm_b:
                images_to_embed.append(norm_b)

    for item in receipts_data or []:
        img_p = item.get("image_path")
        if img_p:
            norm_b = _normalize_image_bytes(img_p)
            if norm_b and norm_b not in images_to_embed:
                images_to_embed.append(norm_b)

    # 4. Handle {{PROOFS}}
    for p in list(doc.paragraphs):
        if "{{PROOFS}}" in p.text:
            if not images_to_embed:
                p.text = "Original receipts verified via Vision OCR."
                if p.runs:
                    p.runs[0].font.name = "Arial"
                    p.runs[0].font.size = Pt(9.5)
                    p.runs[0].font.italic = True
            else:
                p._element.getparent().remove(p._element)

    if images_to_embed:
        for idx, img_bytes in enumerate(images_to_embed):
            try:
                bio = io.BytesIO(img_bytes)
                pil_img = Image.open(bio)
                pil_img = ImageOps.exif_transpose(pil_img)

                if pil_img.mode in ("RGBA", "P"):
                    pil_img = pil_img.convert("RGB")

                buffer = io.BytesIO()
                pil_img.save(buffer, format="JPEG", quality=85)
                buffer.seek(0)

                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.paragraph_format.space_before = Pt(12)
                p_img.paragraph_format.space_after = Pt(14)
                run_img = p_img.add_run()
                run_img.add_picture(buffer, width=Inches(5.0))
            except Exception as e:
                print(f"⚠️ [DOCX Image Embed Error] {e}")

    # Output document
    doc_io = io.BytesIO()
    doc.save(doc_io)
    docx_bytes = doc_io.getvalue()

    if output_path:
        os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
        with open(output_path, "wb") as f:
            f.write(docx_bytes)
        print(f"✅ [DOCX Generator] Successfully saved report to: {output_path}")

    return docx_bytes
